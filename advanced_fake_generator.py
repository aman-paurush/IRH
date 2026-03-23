"""
Advanced fake file generator for realistic honeypot environment.
Creates diverse file types: PDFs, Excel, Word documents, logs, configurations.
"""

import random
import os
from datetime import datetime, timedelta
from pathlib import Path
from io import BytesIO

from faker import Faker

# Try to import optional libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class AdvancedFakeGenerator:
    def __init__(self, base_dir: Path):
        self.base_dir = Path(base_dir)
        self.faker = Faker()

    def generate_pdf_document(self, filepath: Path, doc_type: str = "report") -> bool:
        """Generate a realistic PDF document."""
        if not HAS_REPORTLAB:
            return False

        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            
            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=letter,
                topMargin=72,
                bottomMargin=72,
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#1F4E78',
                spaceAfter=30,
                alignment=1,  # Center
            )
            
            doc_titles = {
                "report": f"Monthly {self.faker.word().title()} Report",
                "financial": f"{self.faker.company()} - Financial Statement",
                "technical": f"Technical Documentation - {self.faker.word().title()}",
                "minutes": f"Meeting Minutes - {datetime.now().strftime('%Y-%m-%d')}",
            }
            
            title = Paragraph(doc_titles.get(doc_type, "Business Document"), title_style)
            elements.append(title)
            elements.append(Spacer(1, 12))
            
            # Add metadata
            metadata_text = f"""
            <b>Document:</b> {self.faker.word()}<br/>
            <b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
            <b>Department:</b> {self.faker.word().title()}<br/>
            <b>Classification:</b> CONFIDENTIAL
            """
            elements.append(Paragraph(metadata_text, styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Add content
            for i in range(3):
                elements.append(Paragraph(f"<b>Section {i+1}: {self.faker.sentence()}</b>", styles['Heading2']))
                for j in range(2):
                    elements.append(Paragraph(self.faker.paragraph(nb_sentences=5), styles['Normal']))
                    elements.append(Spacer(1, 12))
            
            doc.build(elements)
            return True
        except Exception:
            return False

    def generate_excel_spreadsheet(self, filepath: Path, sheet_type: str = "data") -> bool:
        """Generate a realistic Excel spreadsheet."""
        if not HAS_OPENPYXL:
            return False

        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            wb = Workbook()
            ws = wb.active
            ws.title = "Data"
            
            # Add header row
            headers = self._get_headers_for_type(sheet_type)
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col)
                cell.value = header
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            
            # Add data rows
            num_rows = random.randint(50, 200)
            for row in range(2, num_rows + 2):
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=row, column=col)
                    cell.value = self._generate_cell_value(header, sheet_type)
            
            # Adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
            
            wb.save(str(filepath))
            return True
        except Exception:
            return False

    def generate_word_document(self, filepath: Path) -> bool:
        """Generate a realistic Word document."""
        if not HAS_DOCX:
            return False

        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            
            # Add title
            title = doc.add_heading(f"Internal Document - {self.faker.word().title()}", 0)
            title_format = title.runs[0]
            title_format.font.color.rgb = RGBColor(31, 78, 120)
            
            # Add metadata
            doc.add_paragraph(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Department: {self.faker.word().title()}")
            doc.add_paragraph("Classification: CONFIDENTIAL\n")
            
            # Add sections
            for i in range(3):
                doc.add_heading(f"Section {i+1}", level=1)
                doc.add_paragraph(self.faker.paragraph(nb_sentences=5))
                
                # Add a table
                table = doc.add_table(rows=4, cols=3)
                table.style = 'Light Grid Accent 1'
                for j, header in enumerate(['Item', 'Value', 'Status']):
                    table.rows[0].cells[j].text = header
                
                for row_idx in range(1, 4):
                    table.rows[row_idx].cells[0].text = f"Item {row_idx}"
                    table.rows[row_idx].cells[1].text = str(random.randint(1000, 9999))
                    table.rows[row_idx].cells[2].text = random.choice(['Active', 'Pending', 'Complete'])
                
                doc.add_paragraph()
            
            doc.save(str(filepath))
            return True
        except Exception:
            return False

    def generate_log_file(self, filepath: Path) -> bool:
        """Generate a realistic system/application log file."""
        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            
            log_lines = []
            base_time = datetime.now() - timedelta(days=random.randint(1, 7))
            
            log_events = [
                "Application started",
                "User login successful",
                "Database connection established",
                "Configuration loaded",
                "Processing batch job",
                "Backup started",
                "Sync completed",
                "Health check passed",
                "Cache cleared",
                "Performance metrics collected",
            ]
            
            for i in range(random.randint(50, 500)):
                log_time = base_time + timedelta(minutes=i*5)
                event = random.choice(log_events)
                level = random.choice(['INFO', 'DEBUG', 'WARNING', 'ERROR'])
                module = random.choice(['auth', 'database', 'api', 'cache', 'scheduler'])
                log_lines.append(
                    f"[{log_time.strftime('%Y-%m-%d %H:%M:%S')}] [{level}] [{module}] {event}"
                )
            
            filepath.write_text("\n".join(log_lines), encoding="utf-8")
            return True
        except Exception:
            return False

    def generate_config_file(self, filepath: Path) -> bool:
        """Generate a realistic configuration file."""
        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            
            config_content = f"""
# Configuration File Generated {datetime.now().isoformat()}
# DO NOT MODIFY UNLESS AUTHORIZED

[database]
host = db-{random.randint(10, 99)}.internal.company.com
port = {random.choice([5432, 3306, 1433])}
username = {self.faker.user_name()}
password = {"*" * 20}  # ENCRYPTED
database = {self.faker.word()}_{random.randint(1, 9)}
pool_size = {random.randint(5, 20)}

[cache]
enabled = true
backend = redis
host = cache-{random.randint(1, 5)}.internal
expiry_hours = {random.randint(1, 24)}

[logging]
level = INFO
format = json
destination = /var/log/application/main.log
retention_days = 90

[security]
ssl_enabled = true
tls_version = 1.3
certificate_path = /etc/ssl/certs/company.crt
key_path = /etc/ssl/private/company.key

[api]
timeout = {random.randint(30, 120)}
max_retries = {random.randint(3, 10)}
rate_limit_per_minute = {random.randint(100, 1000)}
"""
            filepath.write_text(config_content.strip(), encoding="utf-8")
            return True
        except Exception:
            return False

    def generate_source_code(self, filepath: Path, language: str = "python") -> bool:
        """Generate realistic-looking source code files."""
        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            
            if language == "python":
                code = self._generate_python_source()
            elif language == "sql":
                code = self._generate_sql_source()
            elif language == "json":
                code = self._generate_json_source()
            else:
                code = "# Generated source code\n" + self.faker.paragraph()
            
            filepath.write_text(code, encoding="utf-8")
            return True
        except Exception:
            return False

    def _generate_python_source(self) -> str:
        """Generate realistic Python source code."""
        return f'''#!/usr/bin/env python3
"""
{self.faker.sentence()}
Generated: {datetime.now().isoformat()}
"""

import logging
import json
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

class {self.faker.word().title()}Handler:
    """Handle {self.faker.word()} operations."""
    
    def __init__(self, config_path=None):
        self.config_path = config_path or "config.json"
        self.initialized = False
        self.load_config()
    
    def load_config(self):
        """Load configuration from file."""
        try:
            with open(self.config_path) as f:
                self.config = json.load(f)
            logger.info("Configuration loaded successfully")
            self.initialized = True
        except FileNotFoundError:
            logger.error(f"Config file not found: {{self.config_path}}")
    
    def process_data(self, data):
        """Process input data."""
        if not self.initialized:
            raise RuntimeError("Handler not initialized")
        
        logger.info(f"Processing {{len(data)}} items")
        for item in data:
            self._process_item(item)
    
    def _process_item(self, item):
        """Process a single item."""
        logger.debug(f"Processing item: {{item}}")
        # Processing logic here
        pass

if __name__ == "__main__":
    handler = {self.faker.word().title()}Handler()
    print("Handler initialized")
'''

    def _generate_sql_source(self) -> str:
        """Generate realistic SQL source code."""
        return f'''-- Database schema
-- Generated: {datetime.now().isoformat()}

CREATE TABLE IF NOT EXISTS users (
    id INT PRIMARY KEY AUTO_INCREMENT,
    username VARCHAR(255) NOT NULL UNIQUE,
    email VARCHAR(255) NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    is_active BOOLEAN DEFAULT TRUE
);

CREATE TABLE IF NOT EXISTS transactions (
    id INT PRIMARY KEY AUTO_INCREMENT,
    user_id INT NOT NULL,
    amount DECIMAL(10, 2),
    transaction_type VARCHAR(50),
    status VARCHAR(20),
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (user_id) REFERENCES users(id)
);

CREATE INDEX idx_user_email ON users(email);
CREATE INDEX idx_transaction_user ON transactions(user_id);
CREATE INDEX idx_transaction_date ON transactions(created_at);

-- Sample queries
SELECT COUNT(*) FROM users WHERE is_active = TRUE;
SELECT AVG(amount) FROM transactions WHERE DATE(created_at) = CURDATE();
'''

    def _generate_json_source(self) -> str:
        """Generate realistic JSON source."""
        data = {
            "application": "honeypot-system",
            "version": f"{random.randint(1, 5)}.{random.randint(0, 99)}.{random.randint(0, 99)}",
            "generated": datetime.now().isoformat(),
            "environment": random.choice(["production", "staging", "development"]),
            "services": [
                {"name": "api", "port": 8080, "enabled": True},
                {"name": "database", "port": 5432, "enabled": True},
                {"name": "cache", "port": 6379, "enabled": True},
            ],
            "features": {
                "logging": True,
                "monitoring": True,
                "backup": True,
                "replication": random.choice([True, False]),
            },
        }
        return json.dumps(data, indent=2)

    @staticmethod
    def _get_headers_for_type(sheet_type: str) -> list:
        """Get appropriate headers based on sheet type."""
        headers_map = {
            "financial": ["Date", "Account", "Debit", "Credit", "Balance", "Status"],
            "inventory": ["Item ID", "Product Name", "Quantity", "Unit Cost", "Total Value", "Location"],
            "sales": ["Transaction ID", "Customer", "Product", "Amount", "Date", "Region"],
            "employees": ["Employee ID", "Name", "Department", "Salary", "Hire Date", "Status"],
            "data": ["ID", "Name", "Value", "Category", "Date", "Status"],
        }
        return headers_map.get(sheet_type, headers_map["data"])

    @staticmethod
    def _generate_cell_value(header: str, sheet_type: str):
        """Generate appropriate value for a cell based on header."""
        faker = Faker()
        
        if header in ["Date", "Hire Date"]:
            return faker.date_between(start_date='-2y')
        elif header in ["Amount", "Salary", "Debit", "Credit", "Balance", "Unit Cost", "Total Value"]:
            return round(random.uniform(100, 100000), 2)
        elif header in ["Quantity"]:
            return random.randint(1, 1000)
        elif header in ["ID", "Employee ID", "Transaction ID", "Item ID"]:
            return random.randint(10000, 99999)
        elif header in ["Status"]:
            return random.choice(["Active", "Inactive", "Pending", "Complete", "Failed"])
        elif header in ["Department", "Category"]:
            return random.choice(["Finance", "IT", "HR", "Sales", "Operations"])
        else:
            return faker.name() if "Name" in header else faker.word()

    def generate_diverse_environment(self, base_profile: dict, file_count: int):
        """Generate a diverse fake environment with various file types."""
        user_dir = self.base_dir / base_profile["username"]
        file_types = [
            ("pdf", "report"),
            ("xlsx", "financial"),
            ("docx", None),
            ("log", None),
            ("conf", None),
            ("py", None),
            ("json", None),
            ("sql", None),
        ]

        created = 0
        for i in range(file_count):
            file_ext, subtype = random.choice(file_types)
            folder = random.choice(["Documents", "Spreadsheets", "Projects", "Archive", "Backups", "Logs", "Config"])
            target_dir = user_dir / folder
            
            filename = f"{self.faker.word()}_{random.randint(1000, 9999)}"
            
            if file_ext == "pdf":
                filepath = target_dir / f"{filename}.pdf"
                if self.generate_pdf_document(filepath, subtype or "report"):
                    created += 1
            elif file_ext == "xlsx":
                filepath = target_dir / f"{filename}.xlsx"
                if self.generate_excel_spreadsheet(filepath, subtype or "data"):
                    created += 1
            elif file_ext == "docx":
                filepath = target_dir / f"{filename}.docx"
                if self.generate_word_document(filepath):
                    created += 1
            elif file_ext == "log":
                filepath = target_dir / f"{filename}.log"
                if self.generate_log_file(filepath):
                    created += 1
            elif file_ext == "conf":
                filepath = target_dir / f"{filename}.conf"
                if self.generate_config_file(filepath):
                    created += 1
            elif file_ext == "py":
                filepath = target_dir / f"{filename}.py"
                if self.generate_source_code(filepath, "python"):
                    created += 1
            elif file_ext == "json":
                filepath = target_dir / f"{filename}.json"
                if self.generate_source_code(filepath, "json"):
                    created += 1
            elif file_ext == "sql":
                filepath = target_dir / f"{filename}.sql"
                if self.generate_source_code(filepath, "sql"):
                    created += 1

        return created


if __name__ == "__main__":
    gen = AdvancedFakeGenerator(Path("./test_env"))
    profile = {"username": "test_user", "department": "test"}
    count = gen.generate_diverse_environment(profile, 20)
    print(f"Generated {count} diverse files")
